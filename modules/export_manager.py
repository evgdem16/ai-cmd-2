"""Export dialogs to HTML, PDF, and DOCX."""

import os
import re
from datetime import datetime
from typing import Optional

from modules.logger_setup import get_logger

logger = get_logger()


_HTML_TEMPLATE = """\
<!DOCTYPE html>
<html lang="ru">
<head>
<meta charset="UTF-8">
<meta name="viewport" content="width=device-width, initial-scale=1.0">
<title>{title}</title>
<style>
  body {{ font-family: 'Segoe UI', Arial, sans-serif; max-width: 900px;
          margin: 40px auto; padding: 0 20px; background: #0d1117; color: #c9d1d9; }}
  h1 {{ color: #58a6ff; border-bottom: 1px solid #30363d; padding-bottom: 10px; }}
  .meta {{ color: #8b949e; font-size: .85em; margin-bottom: 30px; }}
  .message {{ margin: 16px 0; padding: 14px 18px; border-radius: 8px; }}
  .user    {{ background: #161b22; border-left: 3px solid #58a6ff; }}
  .assistant {{ background: #0d1117; border-left: 3px solid #3fb950; }}
  .system  {{ background: #1c1c1c; border-left: 3px solid #d29922; font-style: italic; }}
  .role    {{ font-weight: bold; margin-bottom: 6px; color: #8b949e; font-size:.8em;
              text-transform: uppercase; letter-spacing:.08em; }}
  pre      {{ background: #161b22; padding: 12px; border-radius: 6px;
              overflow-x: auto; font-size: .9em; }}
  code     {{ font-family: 'Cascadia Code','Fira Code',monospace; }}
  .ts      {{ color: #484f58; font-size:.75em; float: right; }}
  .stats   {{ margin-top: 40px; padding: 16px; background: #161b22;
              border-radius: 8px; font-size:.85em; }}
  .stats h2 {{ color: #58a6ff; font-size:1em; margin-top:0; }}
  table    {{ border-collapse: collapse; width: 100%; }}
  td, th   {{ padding: 6px 12px; border: 1px solid #30363d; }}
  th       {{ background: #21262d; }}
</style>
</head>
<body>
<h1>💬 {title}</h1>
<div class="meta">
  Created: {created_at} &nbsp;|&nbsp; Updated: {updated_at} &nbsp;|&nbsp;
  Messages: {msg_count}
</div>
{messages_html}
{stats_html}
</body>
</html>
"""


def _escape(text: str) -> str:
    return (text.replace("&", "&amp;").replace("<", "&lt;").replace(">", "&gt;"))


def _md_to_html_basic(text: str) -> str:
    """Very lightweight markdown-to-html: code blocks, bold, italic."""
    # Code blocks
    text = re.sub(
        r"```(\w*)\n(.*?)```",
        lambda m: f"<pre><code class=\"language-{m.group(1)}\">{_escape(m.group(2))}</code></pre>",
        text, flags=re.DOTALL,
    )
    # Inline code
    text = re.sub(r"`([^`]+)`", lambda m: f"<code>{_escape(m.group(1))}</code>", text)
    # Bold
    text = re.sub(r"\*\*(.+?)\*\*", r"<strong>\1</strong>", text)
    # Italic
    text = re.sub(r"\*(.+?)\*", r"<em>\1</em>", text)
    # Newlines
    text = text.replace("\n\n", "</p><p>").replace("\n", "<br>")
    return f"<p>{text}</p>"


class ExportManager:
    def __init__(self, exports_dir: str = "exports"):
        self.exports_dir = exports_dir
        os.makedirs(exports_dir, exist_ok=True)

    def export_html(self, dialog: dict, dest: Optional[str] = None) -> str:
        name = dialog.get("name", "dialog")
        dest = dest or os.path.join(self.exports_dir, f"{name}.html")
        os.makedirs(os.path.dirname(os.path.abspath(dest)), exist_ok=True)

        messages_html = []
        for msg in dialog.get("messages", []):
            role = msg.get("role", "user")
            content = _md_to_html_basic(msg.get("content", ""))
            ts = msg.get("timestamp", "")
            messages_html.append(
                f'<div class="message {role}">'
                f'<div class="role">{role}<span class="ts">{ts}</span></div>'
                f'{content}</div>'
            )

        stats = dialog.get("_stats", {})
        stats_html = (
            '<div class="stats"><h2>📊 Statistics</h2>'
            f'<p>Total messages: {stats.get("total_messages", 0)} &nbsp;|&nbsp; '
            f'Total tokens: {stats.get("total_tokens", 0)} &nbsp;|&nbsp; '
            f'Total response time: {stats.get("total_response_time", 0):.2f}s</p></div>'
        )

        html = _HTML_TEMPLATE.format(
            title=name,
            created_at=dialog.get("created_at", ""),
            updated_at=dialog.get("updated_at", ""),
            msg_count=len(dialog.get("messages", [])),
            messages_html="\n".join(messages_html),
            stats_html=stats_html,
        )
        with open(dest, "w", encoding="utf-8") as fh:
            fh.write(html)
        logger.info(f"Exported HTML: {dest}")
        return dest

    def export_pdf(self, dialog: dict, dest: Optional[str] = None) -> str:
        try:
            from fpdf import FPDF  # type: ignore
        except ImportError:
            raise RuntimeError("fpdf2 not installed. Run: pip install fpdf2")

        name = dialog.get("name", "dialog")
        dest = dest or os.path.join(self.exports_dir, f"{name}.pdf")
        os.makedirs(os.path.dirname(os.path.abspath(dest)), exist_ok=True)

        pdf = FPDF()
        pdf.set_auto_page_break(auto=True, margin=15)
        pdf.add_page()
        pdf.set_font("Helvetica", "B", 16)
        pdf.cell(0, 10, f"Dialog: {name}", ln=True)
        pdf.set_font("Helvetica", size=9)
        pdf.set_text_color(120, 120, 120)
        pdf.cell(0, 6,
                 f"Created: {dialog.get('created_at','')}  "
                 f"Updated: {dialog.get('updated_at','')}", ln=True)
        pdf.ln(4)

        for msg in dialog.get("messages", []):
            role = msg.get("role", "user").upper()
            content = msg.get("content", "")
            ts = msg.get("timestamp", "")

            # Role header
            pdf.set_font("Helvetica", "B", 10)
            pdf.set_text_color(80, 166, 255) if role == "USER" else \
                pdf.set_text_color(63, 185, 80)
            pdf.cell(0, 7, f"[{role}]  {ts}", ln=True)

            # Content
            pdf.set_font("Courier", size=9)
            pdf.set_text_color(50, 50, 50)
            # Strip code fences for PDF plaintext
            plain = re.sub(r"```\w*\n?", "", content)
            pdf.multi_cell(0, 5, plain.encode("latin-1", errors="replace").decode("latin-1"))
            pdf.ln(3)

        stats = dialog.get("_stats", {})
        pdf.set_font("Helvetica", "B", 10)
        pdf.set_text_color(88, 166, 255)
        pdf.cell(0, 8, "Statistics", ln=True)
        pdf.set_font("Helvetica", size=9)
        pdf.set_text_color(80, 80, 80)
        pdf.cell(0, 6,
                 f"Messages: {stats.get('total_messages',0)}  "
                 f"Tokens: {stats.get('total_tokens',0)}  "
                 f"Time: {stats.get('total_response_time',0):.2f}s", ln=True)

        pdf.output(dest)
        logger.info(f"Exported PDF: {dest}")
        return dest

    def export_docx(self, dialog: dict, dest: Optional[str] = None) -> str:
        try:
            from docx import Document  # type: ignore
            from docx.shared import Pt, RGBColor  # type: ignore
        except ImportError:
            raise RuntimeError("python-docx not installed. Run: pip install python-docx")

        name = dialog.get("name", "dialog")
        dest = dest or os.path.join(self.exports_dir, f"{name}.docx")
        os.makedirs(os.path.dirname(os.path.abspath(dest)), exist_ok=True)

        doc = Document()
        doc.add_heading(f"Dialog: {name}", 0)
        meta = doc.add_paragraph()
        meta.add_run(
            f"Created: {dialog.get('created_at','')}  |  "
            f"Updated: {dialog.get('updated_at','')}"
        ).font.color.rgb = RGBColor(0x8B, 0x94, 0x9E)

        for msg in dialog.get("messages", []):
            role = msg.get("role", "user")
            content = msg.get("content", "")
            ts = msg.get("timestamp", "")

            h = doc.add_heading(f"{role.upper()} — {ts}", level=3)
            h.runs[0].font.color.rgb = (
                RGBColor(0x58, 0xA6, 0xFF) if role == "user"
                else RGBColor(0x3F, 0xB9, 0x50)
            )

            # Split code blocks from regular text
            parts = re.split(r"(```\w*\n.*?```)", content, flags=re.DOTALL)
            for part in parts:
                if part.startswith("```"):
                    code = re.sub(r"```\w*\n?", "", part).rstrip("`")
                    p = doc.add_paragraph()
                    run = p.add_run(code)
                    run.font.name = "Courier New"
                    run.font.size = Pt(9)
                else:
                    if part.strip():
                        doc.add_paragraph(part.strip())

        stats = dialog.get("_stats", {})
        doc.add_heading("Statistics", level=2)
        doc.add_paragraph(
            f"Total messages: {stats.get('total_messages',0)}\n"
            f"Total tokens: {stats.get('total_tokens',0)}\n"
            f"Total response time: {stats.get('total_response_time',0):.2f}s"
        )

        doc.save(dest)
        logger.info(f"Exported DOCX: {dest}")
        return dest
